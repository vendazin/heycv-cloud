#!/usr/bin/env python3
"""
heycv.cloud – Backend Server
Handles: job scraping, CV parsing (PDF/DOCX/TXT), Claude API proxy
Run: python3 server.py
"""

import os, re, json, io, traceback, urllib.parse
from flask import Flask, request, jsonify, send_from_directory, Response
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx as python_docx

app = Flask(__name__, static_folder=".")

# ── CORS ──────────────────────────────────────────────────────────────────────
@app.after_request
def add_cors(resp):
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-API-Key"
    resp.headers["Access-Control-Allow-Methods"] = "POST, GET, OPTIONS"
    return resp

@app.route("/", defaults={"path": "index.html"})
@app.route("/<path:path>")
def serve_static(path):
    return send_from_directory(".", path)

@app.route("/api/scrape", methods=["OPTIONS"])
@app.route("/api/optimize", methods=["OPTIONS"])
def preflight():
    return Response("", 204)


# ── SCRAPING ──────────────────────────────────────────────────────────────────

# Rotate through several User-Agent strings to reduce bot detection
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4 Safari/605.1.15",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
]

BASE_HEADERS = {
    "Accept-Language": "de-DE,de;q=0.9,en-US;q=0.8,en;q=0.7",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "DNT": "1",
}

NOISE_TAGS = {"script","style","noscript","iframe","svg","img",
              "header","footer","nav","aside","form","button","input"}

# Priority selectors for extracting job text, platform-specific first
JOB_SELECTORS = [
    # StepStone
    "[data-at='job-ad-text']", "[class*='job-ad-text']",
    "[data-testid='jobad-details-job-description']",
    # Indeed
    "#jobDescriptionText", ".jobsearch-jobDescriptionText", ".job-description__text",
    # LinkedIn
    ".description__text", ".show-more-less-html__markup",
    # Xing
    "[class*='offerDescription']", "[class*='job-description']",
    # Monster / generic
    ".jobAd__content", "[class*='jobDescription']", "[class*='job-description']",
    "[class*='job_description']", "[id*='jobDescription']",
    "[class*='jobDetail']", "[class*='job-detail']",
    "[class*='jobContent']", "[class*='job-content']",
    "[class*='stellenbeschreibung']", "[class*='anzeigentext']",
    "[class*='vacancy-description']", "[class*='position-description']",
    # Workday / Greenhouse / Lever (popular ATS)
    "[data-automation-id='jobPostingDescription']",   # Workday
    "#content", ".content__intro", ".posting-description",  # Lever
    "#app_body", ".job-post-details",                       # Greenhouse-style
    # Generic fallbacks
    "article[class*='job']", "article[class*='offer']",
    "main article", "article", "[role='main']", "main",
]

def fetch_html(url: str) -> str:
    """Fetch HTML with retry logic and multiple UA strings."""
    import random
    headers = {**BASE_HEADERS, "User-Agent": random.choice(USER_AGENTS)}
    session = requests.Session()

    # First attempt – normal
    try:
        r = session.get(url, headers=headers, timeout=14,
                        allow_redirects=True, verify=True)
        if r.status_code < 400:
            return r.text
        # 403 / 429 → try with slightly different headers
        if r.status_code in (403, 429):
            headers["Referer"] = "https://www.google.com/"
            headers["User-Agent"] = random.choice(USER_AGENTS)
            r2 = session.get(url, headers=headers, timeout=14,
                             allow_redirects=True, verify=True)
            if r2.status_code < 400:
                return r2.text
            raise requests.HTTPError(f"HTTP {r2.status_code} – Seite verweigert Zugriff (Bot-Schutz)")
        r.raise_for_status()
        return r.text
    except requests.SSLError:
        r = session.get(url, headers=headers, timeout=14,
                        allow_redirects=True, verify=False)
        r.raise_for_status()
        return r.text


def extract_job_text(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")

    # Remove noise
    for tag in soup(list(NOISE_TAGS)):
        tag.decompose()

    # Try selectors best-first
    for sel in JOB_SELECTORS:
        try:
            el = soup.select_one(sel)
            if el:
                text = el.get_text(separator="\n", strip=True)
                if len(text) > 250:
                    return _clean(text)
        except Exception:
            pass

    # Final fallback: body
    body = soup.find("body")
    raw = (body or soup).get_text(separator="\n", strip=True)
    return _clean(raw)


def _clean(text: str) -> str:
    lines = [l.strip() for l in text.splitlines()]
    lines = [l for l in lines if l and len(l) > 1]
    # Deduplicate adjacent
    out, prev = [], None
    for l in lines:
        if l != prev:
            out.append(l)
        prev = l
    return "\n".join(out)[:14000]


@app.route("/api/scrape", methods=["POST"])
def scrape():
    data = request.get_json(silent=True) or {}
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "Keine URL angegeben."}), 400

    # Basic URL validation
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return jsonify({"error": "Ungültige URL – muss mit http:// oder https:// beginnen."}), 400

    try:
        html = fetch_html(url)
    except requests.HTTPError as e:
        return jsonify({"error": str(e)}), 502
    except requests.ConnectionError:
        return jsonify({"error": "Seite nicht erreichbar – bitte URL prüfen."}), 502
    except requests.Timeout:
        return jsonify({"error": "Zeitüberschreitung beim Laden der Seite."}), 504
    except Exception as e:
        return jsonify({"error": f"Ladefehler: {str(e)[:200]}"}), 500

    text = extract_job_text(html)
    if len(text) < 150:
        return jsonify({"error": "Kein Stellentext gefunden. Die Seite könnte JavaScript erfordern oder den Zugriff blockieren."}), 422

    return jsonify({"text": text, "length": len(text)})


# ── CV PARSING ────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t.strip())
        return "\n\n".join(pages)
    except Exception as e:
        raise ValueError(f"PDF konnte nicht gelesen werden: {e}")


def parse_docx(data: bytes) -> str:
    try:
        doc = python_docx.Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in parts:
                        parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Word-Datei konnte nicht gelesen werden: {e}")


def parse_cv(data: bytes, filename: str) -> str:
    fn = (filename or "").lower()
    if fn.endswith(".pdf"):
        text = parse_pdf(data)
    elif fn.endswith((".docx", ".doc")):
        text = parse_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")
    text = text.strip()
    if len(text) < 40:
        raise ValueError("Lebenslauf konnte nicht ausgelesen werden oder ist leer.")
    return text[:16000]


# ── PROMPT BUILDER ────────────────────────────────────────────────────────────

def build_prompt(cv: str, job: str, opts: dict) -> str:
    opt_lines = []
    if opts.get("ats"):
        opt_lines.append("• ATS-Optimierung: Klare Struktur, ATS-konforme Sprache, keine Sonderzeichen in Überschriften")
    if opts.get("keywords"):
        opt_lines.append("• Keyword-Matching: Exakte Formulierungen und Fachbegriffe aus der Stellenanzeige einbauen")
    if opts.get("summary"):
        opt_lines.append("• Profil-Summary: 3-4 prägnante Sätze direkt auf diese Stelle zugeschnitten, ganz oben")
    if opts.get("cover"):
        opt_lines.append("• Anschreiben-Tipps: 3-5 konkrete, stellenspezifische Tipps am Ende des Outputs")
    if not opt_lines:
        opt_lines.append("• Allgemeine Optimierung: Relevanz für diese Stelle maximieren")

    return f"""Du bist ein erstklassiger Karriereberater und CV-Spezialist mit 15 Jahren Erfahrung in DACH-Unternehmen, HR und Recruiting.

═══ DEINE AUFGABE ═══
Optimiere den Lebenslauf präzise und gezielt für diese Stellenanzeige. Halte alle Fakten bei – erfinde nichts.

═══ STELLENANZEIGE ═══
{job}

═══ ORIGINAL-LEBENSLAUF ═══
{cv}

═══ GEWÜNSCHTE OPTIMIERUNGEN ═══
{chr(10).join(opt_lines)}

═══ VORGEHEN ═══
1. Analysiere: Welche konkreten Skills, Erfahrungen, Soft-Skills und Keywords werden gesucht?
2. Passe den Lebenslauf an: Hebe relevante Stationen hervor, nutze die Sprache der Ausschreibung
3. Integriere Keywords organisch – niemals aufgesetzt oder wiederholend
4. Starke Verben am Anfang jeder Bullet-Point (Entwickelte, Leitete, Optimierte, Steigerte…)
5. Quantifiziere Erfolge wo möglich (%, Zahlen, Zeiträume)
6. Behalte exakt die Sprache bei, in der die Stellenanzeige verfasst ist (DE/EN)

═══ AUSGABEFORMAT (exakt so, keine Abweichungen) ═══

---MATCH_SCORE---
[Nur die Zahl, z.B. 88]

---ATS_SCORE---
[Nur die Zahl, z.B. 93]

---KEYWORDS_ADDED---
[Nur die Zahl der neu integrierten Keywords, z.B. 11]

---OPTIMIERTER_LEBENSLAUF---
[Vollständiger, druckfertiger Lebenslauf – mit Profil-Summary oben, dann Berufserfahrung, Ausbildung, Skills. Klar strukturiert.]

---VERBESSERUNGEN---
[Bullet-Liste der wichtigsten Änderungen, je Punkt mit • beginnend. Was wurde konkret verändert und warum?]

---TIPPS---
[3-5 konkrete Handlungsempfehlungen für diese Bewerbung, je mit • beginnend. Z.B. worauf im Anschreiben eingehen, was beim Interview wichtig ist.]"""


# ── OPTIMIZE ──────────────────────────────────────────────────────────────────

@app.route("/api/optimize", methods=["POST"])
def optimize():
    api_key = request.headers.get("X-API-Key", "").strip()
    if not api_key:
        return jsonify({"error": "Kein API-Key angegeben. Bitte trage deinen Anthropic API-Key in den Einstellungen ein."}), 401

    cv_file  = request.files.get("cv")
    job_text = (request.form.get("jobText") or "").strip()
    job_url  = (request.form.get("jobUrl")  or "").strip()
    opts_raw = request.form.get("options", "{}")

    try:
        opts = json.loads(opts_raw)
    except Exception:
        opts = {}

    if not cv_file:
        return jsonify({"error": "Kein Lebenslauf hochgeladen."}), 400
    if not job_text and not job_url:
        return jsonify({"error": "Kein Stellenanzeigetext und keine URL angegeben."}), 400

    # Parse CV
    try:
        cv_bytes = cv_file.read()
        cv_text  = parse_cv(cv_bytes, cv_file.filename or "cv.txt")
    except ValueError as e:
        return jsonify({"error": str(e)}), 422

    # Fallback scrape if jobText empty
    if not job_text and job_url:
        try:
            html = fetch_html(job_url)
            job_text = extract_job_text(html)
        except Exception as e:
            return jsonify({"error": f"Stelle konnte nicht geladen werden: {e}"}), 502

    if len(job_text) < 100:
        return jsonify({"error": "Stellentext zu kurz oder nicht lesbar. Bitte URL prüfen."}), 422

    # Call Claude
    prompt = build_prompt(cv_text, job_text, opts)
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": "claude-sonnet-4-5",
                "max_tokens": 4096,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=120,
        )
        resp.raise_for_status()
        data   = resp.json()
        output = "".join(
            b.get("text", "") for b in data.get("content", []) if b.get("type") == "text"
        )
        return jsonify({"result": output})

    except requests.HTTPError as e:
        try:
            body = resp.json()
            msg  = body.get("error", {}).get("message", str(e))
        except Exception:
            msg = str(e)
        if "401" in str(e) or "authentication" in msg.lower() or "api_key" in msg.lower():
            return jsonify({"error": "Ungültiger API-Key. Bitte prüfe deinen Anthropic API-Key unter console.anthropic.com."}), 401
        return jsonify({"error": f"Claude API-Fehler: {msg}"}), 502
    except requests.Timeout:
        return jsonify({"error": "Claude-Antwort hat zu lange gebraucht. Bitte erneut versuchen."}), 504
    except Exception as e:
        return jsonify({"error": f"Serverfehler: {str(e)[:300]}"}), 500


# ── RUN ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"\n🚀 heycv.cloud läuft auf  http://localhost:{port}")
    print(f"   Web-App:               http://localhost:{port}/index.html\n")
    app.run(host="0.0.0.0", port=port, debug=False)
